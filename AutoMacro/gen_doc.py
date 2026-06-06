from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

IMG_DIR = r'c:\Users\Ye\Documents\叶汇力\全能脚本精灵\说明截图'
IMG1 = IMG_DIR + r'\ScreenShot_2026-03-29_211850_908.png'  # 主界面
IMG2 = IMG_DIR + r'\ScreenShot_2026-03-29_211925_380.png'  # 新建方案
IMG3 = IMG_DIR + r'\ScreenShot_2026-03-29_212003_069.png'  # 录制中
IMG4 = IMG_DIR + r'\ScreenShot_2026-03-29_212019_678.png'  # 插入变量
IMG5 = IMG_DIR + r'\ScreenShot_2026-03-29_212045_094.png'  # 变量表

def add_img(doc, path, width=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        run.font.italic = True

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('小叶专用PC脚本')
run.font.size = Pt(32)
run.font.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('多方案键鼠录制与自动化回放工具')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(100, 100, 100)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('使用说明书')
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(130, 130, 130)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading('目录', level=1)
toc = ['一、软件简介', '二、安装与启动', '三、创建方案', '四、录制操作',
       '五、变量表配置', '六、回放执行', '七、方案管理', '八、快捷键配置',
       '九、跨电脑使用', '十、文件结构']
for item in toc:
    doc.add_paragraph(item).paragraph_format.space_after = Pt(6)
doc.add_page_break()

# ===== 一、软件简介 =====
doc.add_heading('一、软件简介', level=1)
doc.add_paragraph(
    '小叶专用PC脚本是一款多方案键鼠录制与自动化回放工具。'
    '支持录制鼠标移动、点击、滚轮及键盘操作，并可在回放时自动注入动态变量'
    '（如卡密、账号等），实现批量注册、环境初始化等自动化任务。')
doc.add_paragraph('')
doc.add_paragraph('主要特性：')
doc.add_paragraph('  - 多方案独立管理，互不干扰')
doc.add_paragraph('  - 录制时按 F9 插入变量标记，回放自动从变量表逐行粘贴')
doc.add_paragraph('  - 支持循环回放，配合变量表实现批量操作')
doc.add_paragraph('  - 方案可导出为 .zip，方便跨电脑搬运')
doc.add_paragraph('  - 全局快捷键可自定义配置')

# ===== 二、安装与启动 =====
doc.add_heading('二、安装与启动', level=1)
doc.add_paragraph('1. 将 AutoMacro.exe 放到任意文件夹中')
doc.add_paragraph('2. 双击运行即可，无需安装，无需 .NET 运行时')
doc.add_paragraph('3. 首次运行会自动在程序目录下创建：')
doc.add_paragraph('   - Profiles/ 文件夹（存放方案数据）')
doc.add_paragraph('   - hotkey_settings.json（快捷键配置文件）')
doc.add_paragraph('')
doc.add_paragraph('启动后界面如下：')
add_img(doc, IMG1, 6.0, '图1：软件主界面')

# ===== 三、创建方案 =====
doc.add_heading('三、创建方案', level=1)
doc.add_paragraph(
    '使用前需要先创建一个"方案"。每个方案是一组独立的录制脚本和变量数据，互不影响。')
doc.add_paragraph('')
doc.add_paragraph('操作步骤：')
doc.add_paragraph('1. 点击左侧栏的蓝色 "+" 按钮')
doc.add_paragraph('2. 在弹出的对话框中输入方案名称（如"账号注册"）')
doc.add_paragraph('3. 点击"创建"按钮')
doc.add_paragraph('')
add_img(doc, IMG2, 6.0, '图2：点击 "+" 按钮后弹出新建方案对话框')
doc.add_paragraph('')
doc.add_paragraph('创建成功后，左侧导航树会显示：')
doc.add_paragraph('  📁 方案名称')
doc.add_paragraph('    📜 动作脚本 — 点击查看/编辑录制的动作列表')
doc.add_paragraph('    📊 变量表 — 点击查看/编辑变量数据')

# ===== 四、录制操作 =====
doc.add_heading('四、录制操作', level=1)

doc.add_heading('4.1 开始录制', level=2)
doc.add_paragraph('1. 在左侧下拉框中选中目标方案')
doc.add_paragraph('2. 点击底部红色"开始录制 (F10)"按钮，或直接按 F10 快捷键')
doc.add_paragraph('3. 底部状态灯变红，显示"录制中..."')
doc.add_paragraph('4. 此时正常操作鼠标和键盘，所有动作都会被精确记录（含时间间隔）')

doc.add_heading('4.2 暂停/继续录制', level=2)
doc.add_paragraph('录制过程中按 F8 可暂停，暂停期间的任何操作不会被记录。再按 F8 继续。')

doc.add_heading('4.3 插入变量标记（核心功能）', level=2)
doc.add_paragraph(
    '当录制过程中遇到需要输入动态内容的位置（如输入框需要填入卡密），可以插入变量标记：')
doc.add_paragraph('')
doc.add_paragraph('1. 将鼠标光标定位到目标输入框')
doc.add_paragraph('2. 按 F9，录制自动暂停并弹出变量选择对话框')
doc.add_paragraph('3. 在下拉框中选择已有的变量列名，或直接输入新变量名（如"卡密"）')
doc.add_paragraph('4. 点击"插入"按钮')
doc.add_paragraph('5. 动作列表中会新增一条"粘贴变量: 卡密"记录，录制自动继续')
doc.add_paragraph('')
add_img(doc, IMG4, 4.0, '图3：按 F9 弹出的变量插入对话框')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('提示：')
run.font.bold = True
p.add_run('如果输入的变量名在变量表中不存在，系统会自动在变量表中创建对应的新列。')

doc.add_heading('4.4 停止录制', level=2)
doc.add_paragraph('点击底部"停止录制 (F10)"按钮，或再次按 F10。')
doc.add_paragraph('录制的所有动作会自动保存到方案的 record.json 文件中，并实时显示在动作列表中：')
doc.add_paragraph('')
add_img(doc, IMG3, 6.0, '图4：录制完成后的动作列表，记录了鼠标移动轨迹和时间间隔')

# ===== 五、变量表配置 =====
doc.add_heading('五、变量表配置', level=1)
doc.add_paragraph('点击左侧导航树中的"变量表"，切换到变量编辑器界面。')
doc.add_paragraph('')
add_img(doc, IMG5, 6.0, '图5：变量表编辑器界面')

doc.add_heading('5.1 添加变量列', level=2)
doc.add_paragraph('1. 在顶部"添加列"输入框中输入变量列名（如"卡密"、"账号"）')
doc.add_paragraph('2. 点击"添加列"按钮，表格中会出现新的一列')

doc.add_heading('5.2 填写变量数据', level=2)
doc.add_paragraph('1. 点击"添加行"按钮，添加数据行')
doc.add_paragraph('2. 双击单元格即可直接编辑内容')
doc.add_paragraph('3. 每一行对应回放时的一次粘贴：')
doc.add_paragraph('   - 第 1 行的值 = 第 1 次遇到该变量时粘贴的内容')
doc.add_paragraph('   - 第 2 行的值 = 第 2 次遇到该变量时粘贴的内容')
doc.add_paragraph('   - 以此类推...')

doc.add_heading('5.3 使用示例', level=2)
doc.add_paragraph('假设需要批量注册 3 个账号，变量表可以这样填：')
doc.add_paragraph('')
table = doc.add_table(rows=4, cols=2, style='Table Grid')
for i, (a, b) in enumerate([('卡密', '账号'),
                             ('ABC-001', 'user1'),
                             ('ABC-002', 'user2'),
                             ('ABC-003', 'user3')]):
    table.cell(i, 0).text = a
    table.cell(i, 1).text = b
for cell in table.rows[0].cells:
    for para in cell.paragraphs:
        for r in para.runs:
            r.font.bold = True

doc.add_paragraph('')
doc.add_paragraph(
    '回放时，脚本每次遇到"粘贴变量: 卡密"就从"卡密"列依次取值：'
    '第1次取 ABC-001，第2次取 ABC-002，第3次取 ABC-003。')

doc.add_heading('5.4 列管理', level=2)
doc.add_paragraph('  - 删除列：在"选择列"下拉框中选中目标列，点击红色"删除列"按钮')
doc.add_paragraph('  - 重命名列：选中后点击"重命名列"按钮，在弹窗中输入新列名')
doc.add_paragraph('  - 清空全部：点击"清空全部"一键删除所有列和数据')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('重要：')
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
p.add_run('编辑完成后务必点击右上角"保存 CSV"按钮，否则数据不会写入文件。')

# ===== 六、回放执行 =====
doc.add_heading('六、回放执行', level=1)

doc.add_heading('6.1 设置循环次数', level=2)
doc.add_paragraph('在底部"循环次数"输入框中设置回放轮次：')
doc.add_paragraph('  - 1 = 执行一次脚本')
doc.add_paragraph('  - 3 = 连续执行三次（每次从变量表依次读取数据）')
doc.add_paragraph('  - 0 = 无限循环（直到手动停止或变量数据用完）')

doc.add_heading('6.2 开始回放', level=2)
doc.add_paragraph('1. 确保已选中方案，且动作列表中有录制内容')
doc.add_paragraph('2. 点击底部蓝色"开始回放 (F12)"按钮，或按 F12')
doc.add_paragraph('3. 状态灯变蓝，程序开始按照录制的时间轴自动模拟键鼠操作')
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('注意：')
run.font.bold = True
p.add_run('回放期间请不要移动鼠标或操作键盘，以免干扰自动化流程。')

doc.add_heading('6.3 紧急停止', level=2)
doc.add_paragraph('回放或录制过程中，随时可以：')
doc.add_paragraph('  - 点击底部橙色"紧急停止"按钮')
doc.add_paragraph('  - 或按 F11 快捷键')
doc.add_paragraph('程序会立即中断当前操作。')

# ===== 七、方案管理 =====
doc.add_heading('七、方案管理', level=1)
doc.add_paragraph('左侧栏提供完整的方案管理功能：')
doc.add_paragraph('')
table2 = doc.add_table(rows=6, cols=2, style='Table Grid')
rows_data = [('操作', '说明'),
             ('新建方案', '点击蓝色 "+" 按钮，输入名称后创建'),
             ('重命名', '点击笔形按钮，在弹窗中修改名称'),
             ('删除方案', '点击红色垃圾桶按钮，确认后彻底删除'),
             ('导出方案', '点击"导出方案"按钮，保存为 .zip 压缩包'),
             ('导入方案', '点击"导入方案"按钮，选择 .zip 文件导入')]
for i, (a, b) in enumerate(rows_data):
    table2.cell(i, 0).text = a
    table2.cell(i, 1).text = b
for cell in table2.rows[0].cells:
    for para in cell.paragraphs:
        for r in para.runs:
            r.font.bold = True

# ===== 八、快捷键配置 =====
doc.add_heading('八、快捷键配置', level=1)
doc.add_paragraph('软件默认快捷键如下：')
doc.add_paragraph('')
hotkeys = [('功能', '默认快捷键'),
           ('开始/停止录制', 'F10'),
           ('暂停/继续录制', 'F8'),
           ('插入变量', 'F9'),
           ('紧急停止', 'F11'),
           ('开始/停止回放', 'F12')]
table3 = doc.add_table(rows=len(hotkeys), cols=2, style='Table Grid')
for i, (a, b) in enumerate(hotkeys):
    table3.cell(i, 0).text = a
    table3.cell(i, 1).text = b
for cell in table3.rows[0].cells:
    for para in cell.paragraphs:
        for r in para.runs:
            r.font.bold = True

doc.add_paragraph('')
doc.add_paragraph(
    '如需自定义快捷键，编辑程序目录下的 hotkey_settings.json 文件（用记事本打开即可）：')
doc.add_paragraph('')
code = '{\n  "StartStopRecording": "F10",\n  "PauseRecording": "F8",\n  "InsertVariable": "F9",\n  "EmergencyStop": "F11",\n  "StartPlayback": "F12"\n}'
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('')
doc.add_paragraph('支持的键名：F1 ~ F12、Pause、ScrollLock、Insert、Home、End、PageUp、PageDown')
doc.add_paragraph('修改后需要重启软件才会生效。')

# ===== 九、跨电脑使用 =====
doc.add_heading('九、跨电脑使用', level=1)

doc.add_heading('前提条件', level=2)
p = doc.add_paragraph()
run = p.add_run('重要：')
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
p.add_run(
    '本软件使用绝对屏幕坐标录制鼠标操作。在不同电脑上回放时，'
    '必须保证两台电脑的屏幕分辨率和系统缩放比例完全一致，否则鼠标点击位置会出现偏移。')
doc.add_paragraph('')
doc.add_paragraph('例如：A 电脑 2560x1440 + 100% 缩放录制 → B 电脑也必须是 2560x1440 + 100% 缩放')

doc.add_heading('搬运步骤', level=2)
doc.add_paragraph('1. 在 A 电脑上选中目标方案，点击左侧"导出方案"按钮')
doc.add_paragraph('2. 选择保存位置，生成 .zip 压缩包')
doc.add_paragraph('3. 将 .zip 文件通过 U 盘、微信、网盘等方式传到 B 电脑')
doc.add_paragraph('4. 在 B 电脑的小叶专用PC脚本中，点击"导入方案"按钮')
doc.add_paragraph('5. 选择 .zip 文件，方案自动导入并显示在左侧列表中')

# ===== 十、文件结构 =====
doc.add_heading('十、文件结构', level=1)
doc.add_paragraph('程序运行后的文件目录结构如下：')
doc.add_paragraph('')
structure = ('AutoMacro.exe                  主程序\n'
             'hotkey_settings.json           快捷键配置\n'
             'Profiles/                      方案数据目录\n'
             '  ├── 方案一/\n'
             '  │   ├── record.json          动作脚本（鼠标键盘操作记录）\n'
             '  │   └── variables.csv        变量表（卡密、账号等数据）\n'
             '  └── 方案二/\n'
             '      ├── record.json\n'
             '      └── variables.csv')
p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# 保存
out = r'c:\Users\Ye\Documents\叶汇力\全能脚本精灵\AutoMacro\publish\小叶专用PC脚本_使用说明.docx'
doc.save(out)
print(f'Done: {out}')
